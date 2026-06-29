from datetime import date
from decimal import Decimal

import pytest
from docx import Document
from io import BytesIO

from services.docx_generator import _format_date, _format_item_price, _render, generate_single, generate_two, generate_three


@pytest.fixture
def single_employee():
    """Фикстура, которая возвращает словарь с данными одного сотрудника для тестов."""
    return {
        "full_name": "Иванов Иван Иванович",
        "position": "Механик",
        "contract_date": date(2024, 1, 10),
        "contract_number": "1-aa",
        "document_type": "Паспорт РФ",
        "id_series": "1234",
        "id_number": "456789",
        "id_issued_date": date(2000, 1, 1),
        "issued_by": "Подразделением",
        "address": "г. Екатеринбург, ул. Ленина, д. 1"
    }


@pytest.fixture
def second_employee():
    """Фикстура, которая возвращает словарь с данными второго сотрудника для тестов."""
    return {
        "full_name": "Петров Петр Петрович",
        "position": "Маляр",
        "contract_date": date(2024, 11, 1),
        "contract_number": "2-bb",
        "document_type": "ВНЖ",
        "id_series": "0000",
        "id_number": "123456",
        "id_issued_date": date(2020, 2, 2),
        "issued_by": "Кем то",
        "address": "г. Екатеринбург, ул. Малышева, д. 2"
    }


@pytest.fixture
def third_employee():
    """Фикстура, которая возвращает словарь с данными третьего сотрудника для тестов."""
    return {
        "full_name": "Сидоров Сидор Сидорович",
        "position": "Слесарь",
        "contract_date": date(2023, 5, 15),
        "contract_number": "3-cc",
        "document_type": "Паспорт РФ",
        "id_series": "5678",
        "id_number": "987654",
        "id_issued_date": date(2015, 6, 20),
        "issued_by": "Третьим отделением",
        "address": "г. Екатеринбург, ул. Мира, д. 3"
    }


@pytest.fixture
def items():
    """Фикстура, которая возвращает список словарей с данными об инвентаре для тестов."""
    return [
        {"tool_name": "Вороток шарнирный", "tool_code": "ЦБ000001111", "quantity": 2, "price": Decimal("2500.00")},
        {"tool_name": "Набор ключей", "tool_code": "БШ-000002222", "quantity": 1, "price": Decimal("1000.00")}
    ]


def test_format_date(single_employee):
    """Тестирует форматирование даты в строку формата '10 января 2024'."""
    d = single_employee["contract_date"]
    assert _format_date(d) == "10 января 2024"


def test_format_item_price(items):
    """Тестирует форматирование цены в строку формата '2 500,00'."""
    price = items[0]["price"]
    assert _format_item_price(price) == "2 500,00"


def test_render(single_employee):
    """
    Тестирует рендеринг шаблона с контекстом.
    Проверяет, что результат - это байты, которые можно открыть как документ.
    """
    context = {**single_employee, "items": []}
    result = _render("templates/single_employee.example.docx", context)
    doc = Document(BytesIO(result))
    # Проверяем, что документ успешно открыт
    assert doc is not None
    # Проверяем, что результат - это байты
    assert isinstance(result, bytes)


def test_generate_single(single_employee, items):
    """
    Тестирует генерацию документа для одного сотрудника.
    Проверяет, что в документе содержатся все поля сотрудника и отформатированные даты.
    Также проверяет, что таблица с инвентарем заполнилась и содержит название инструмента и отформатированную цену.
    """
    result = generate_single(single_employee, items)
    doc = Document(BytesIO(result))

    # Проверяем, что документ содержит все необходимые данные сотрудника
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Иванов Иван Иванович" in full_text
    assert "Механик" in full_text
    assert "10 января 2024" in full_text
    assert "1-aa" in full_text
    assert "Паспорт РФ" in full_text
    assert "1234" in full_text
    assert "456789" in full_text
    assert "1 января 2000" in full_text
    assert "Подразделением" in full_text
    assert "г. Екатеринбург, ул. Ленина, д. 1" in full_text

    # Проверяем, что таблица с инвентарем заполнилась и содержит название инструмента и отформатированную цену
    for table in doc.tables:
        table_text = "\n".join(cell.text for row in table.rows for cell in row.cells)
    assert "Вороток шарнирный" in table_text
    assert "Набор ключей" in table_text
    assert "2 500,00" in table_text
    assert "1 000,00" in table_text


def test_generate_two(single_employee, second_employee, items):
    """
    Тестирует генерацию документа для двух сотрудников.
    Проверяет, что в документе содержатся все поля всех сотрудников и отформатированные даты.
    Также проверяет, что таблица с инвентарем заполнилась и содержит название инструмента и отформатированную цену.
    """
    result = generate_two(single_employee, second_employee, items)
    doc = Document(BytesIO(result))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Иванов Иван Иванович" in full_text
    assert "Механик" in full_text
    assert "10 января 2024" in full_text
    assert "1-aa" in full_text
    assert "Паспорт РФ" in full_text
    assert "1234" in full_text
    assert "456789" in full_text
    assert "1 января 2000" in full_text
    assert "Подразделением" in full_text
    assert "г. Екатеринбург, ул. Ленина, д. 1" in full_text

    assert "Петров Петр Петрович" in full_text
    assert "Маляр" in full_text
    assert "1 ноября 2024" in full_text
    assert "2-bb" in full_text
    assert "ВНЖ" in full_text
    assert "0000" in full_text
    assert "123456" in full_text
    assert "2 февраля 2020" in full_text
    assert "Кем то" in full_text
    assert "г. Екатеринбург, ул. Малышева, д. 2" in full_text

    for table in doc.tables:
        table_text = "\n".join(cell.text for row in table.rows for cell in row.cells)
    assert "Вороток шарнирный" in table_text
    assert "Набор ключей" in table_text
    assert "2 500,00" in table_text
    assert "1 000,00" in table_text


def test_generate_three_smoke(single_employee, second_employee, third_employee, items):
    """
    Smoke-тест генерации документа для трёх сотрудников.
    Не дублирует детальную проверку каждого поля, т.к. это уже проверено в тестах для одного и двух сотрудников.
    Ловит опечатку в плейсхолдерах внутри шаблона three_employees.docx.
    Поверяет, что документ рендерится без ошибок и ФИО всех трех сотрудников попали в тест.
    """
    result = generate_three(single_employee, second_employee, third_employee, items)
    doc = Document(BytesIO(result))

    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Иванов Иван Иванович" in full_text
    assert "Петров Петр Петрович" in full_text
    assert "Сидоров Сидор Сидорович" in full_text

    for table in doc.tables:
        table_text = "\n".join(cell.text for row in table.rows for cell in row.cells)
    assert "Вороток шарнирный" in table_text
    assert "Набор ключей" in table_text